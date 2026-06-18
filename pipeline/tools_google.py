# Created: 2026-05-20
# Purpose: Google API tools — Gmail, Calendar, Drive (split from tools.py)
# Dependencies: pipeline.auth.google, stdlib, openpyxl(optional), pdfminer/pypdf(optional)

from __future__ import annotations

import json
import urllib.parse
import urllib.request
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from pipeline.auth.google import get_access_token as _google_token


# ── Internal helpers ──────────────────────────────────────────────────────────

def _gapi(path: str, account: str = "", params: dict | None = None,
          method: str = "GET", body: dict | None = None) -> dict:
    """Google API helper with readable OAuth/API errors.

    Previously urllib.HTTPError bubbled up as just `HTTP Error 400`, which hid
    actionable Google payloads such as `invalid_grant` (revoked/expired refresh
    token). Tool callers now receive a compact RuntimeError including the body.
    """
    import urllib.error

    token = _google_token(account)
    if not token:
        raise RuntimeError(
            "Google OAuth access token unavailable. "
            "설정 → 워크스페이스에서 Google 계정을 다시 연결하세요."
        )

    url = path if path.startswith("http") else f"https://{path}"
    if params:
        parts = []
        for k, v in params.items():
            if v is None:
                continue
            if isinstance(v, list):
                for item in v:
                    parts.append((k, item))
            else:
                parts.append((k, v))
        if parts:
            url += "?" + urllib.parse.urlencode(parts)
    data = json.dumps(body).encode() if body else None
    headers = {"Authorization": f"Bearer {token}"}
    if data:
        headers["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        import ssl
        try:
            import certifi
            _ssl_ctx = ssl.create_default_context(cafile=certifi.where())
        except ImportError:
            _ssl_ctx = ssl.create_default_context()
        with urllib.request.urlopen(req, timeout=20, context=_ssl_ctx) as r:
            raw = r.read()
            return json.loads(raw) if raw else {}
    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Google API HTTP {e.code}: {err_body}") from e


# ── Gmail ──────────────────────────────────────────────────────────────────────

def gmail_search(query: str, max_results: int = 10, account: str = "") -> list[dict]:
    base = "gmail.googleapis.com/gmail/v1/users/me"
    data = _gapi(f"{base}/messages", account=account,
                 params={"q": query, "maxResults": max_results})
    messages = data.get("messages", [])
    result = []
    for m in messages:
        detail = _gapi(f"{base}/messages/{m['id']}", account=account,
                       params={"format": "metadata",
                               "metadataHeaders": ["From", "Subject", "Date"]})
        headers = {h["name"]: h["value"] for h in detail.get("payload", {}).get("headers", [])}
        result.append({
            "id": m["id"],
            "threadId": m["threadId"],
            "snippet": detail.get("snippet", ""),
            "from": headers.get("From", ""),
            "subject": headers.get("Subject", ""),
            "date": headers.get("Date", ""),
        })
    return result


def _html_to_text(html: str) -> str:
    import re
    from html import unescape
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html,
                  flags=re.IGNORECASE | re.DOTALL)
    html = re.sub(r"<(br|/p|/div|/tr|/li|/h[1-6])[^>]*>", "\n", html,
                  flags=re.IGNORECASE)
    html = re.sub(r"<[^>]+>", "", html)
    text = unescape(html)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def gmail_read(message_id: str, account: str = "", max_chars: int = 20000) -> dict:
    import base64
    base = "gmail.googleapis.com/gmail/v1/users/me"
    detail = _gapi(f"{base}/messages/{message_id}", account=account,
                   params={"format": "full"})
    headers = {h["name"]: h["value"]
               for h in detail.get("payload", {}).get("headers", [])}

    def _decode(part: dict) -> str:
        data = part.get("body", {}).get("data", "")
        if not data:
            return ""
        return base64.urlsafe_b64decode(data + "==").decode("utf-8", errors="replace")

    def _collect(payload: dict, acc: dict):
        mime = payload.get("mimeType", "")
        if mime == "text/plain" and not acc.get("plain"):
            acc["plain"] = _decode(payload)
        elif mime == "text/html" and not acc.get("html"):
            acc["html"] = _decode(payload)
        for part in payload.get("parts", []):
            _collect(part, acc)

    acc: dict = {}
    _collect(detail.get("payload", {}), acc)

    if acc.get("plain", "").strip():
        body = acc["plain"]
    elif acc.get("html", "").strip():
        body = _html_to_text(acc["html"])
    else:
        body = detail.get("snippet", "")

    body = body.strip()
    truncated = len(body) > max_chars
    return {
        "subject": headers.get("Subject", ""),
        "from": headers.get("From", ""),
        "date": headers.get("Date", ""),
        "body": body[:max_chars],
        "truncated": truncated,
    }


def gmail_list_attachments(message_id: str, account: str = "") -> list[dict]:
    """Returns the list of attachments for a message. Each entry: filename, mime_type, size, attachment_id."""
    import base64
    base = "gmail.googleapis.com/gmail/v1/users/me"
    detail = _gapi(f"{base}/messages/{message_id}", account=account,
                   params={"format": "full"})

    attachments: list[dict] = []

    def _walk(part: dict):
        filename = part.get("filename", "")
        body = part.get("body", {})
        att_id = body.get("attachmentId", "")
        if filename and att_id:
            attachments.append({
                "filename": filename,
                "mime_type": part.get("mimeType", ""),
                "size": body.get("size", 0),
                "attachment_id": att_id,
            })
        for sub in part.get("parts", []):
            _walk(sub)

    _walk(detail.get("payload", {}))
    return attachments


def gmail_download_attachment(
    message_id: str,
    attachment_id: str,
    save_path: str,
    account: str = "",
) -> dict:
    """
    Downloads a Gmail attachment and saves it to a local file.
    save_path: absolute path to save to (parent directory is created if missing).
    Returns: {"saved": True, "path": ..., "size": ...}
    """
    import base64
    base = "gmail.googleapis.com/gmail/v1/users/me"
    data = _gapi(
        f"{base}/messages/{message_id}/attachments/{attachment_id}",
        account=account,
    )
    raw = base64.urlsafe_b64decode(data["data"] + "==")

    dest = Path(save_path)
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(raw)
    return {"saved": True, "path": str(dest), "size": len(raw)}


def _md_to_html(text: str) -> str:
    """
    Markdown → HTML conversion (no external dependencies).
    Briefing-mail level: headings, bold/italic, code, tables, lists, horizontal rules.
    """
    import re as _re

    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        # Headings
        if line.startswith("### "):
            out.append(f"<h3>{line[4:]}</h3>"); i += 1; continue
        if line.startswith("## "):
            out.append(f"<h2>{line[3:]}</h2>"); i += 1; continue
        if line.startswith("# "):
            out.append(f"<h1>{line[2:]}</h1>"); i += 1; continue
        # Horizontal rule
        if _re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", line.strip()):
            out.append("<hr>"); i += 1; continue
        # Table: next line matches |---|--- pattern
        if "|" in line and i + 1 < len(lines) and _re.match(r"^\|?[\s\-\|:]+\|", lines[i + 1]):
            cols = [c.strip() for c in line.strip("|").split("|")]
            out.append('<table style="border-collapse:collapse;width:100%;margin:8px 0">')
            out.append("<thead><tr>" + "".join(
                f'<th style="border:1px solid #444;padding:6px 10px;background:#1e1e2e;text-align:left">{c}</th>'
                for c in cols
            ) + "</tr></thead><tbody>")
            i += 2  # Skip separator row
            while i < len(lines) and "|" in lines[i]:
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                out.append("<tr>" + "".join(
                    f'<td style="border:1px solid #333;padding:5px 10px">{c}</td>'
                    for c in cells
                ) + "</tr>")
                i += 1
            out.append("</tbody></table>")
            continue
        # List item
        if _re.match(r"^[-*]\s", line):
            out.append(f'<li style="margin:2px 0">{line[2:]}</li>'); i += 1; continue
        # Blank line
        if not line.strip():
            out.append("<br>"); i += 1; continue
        # Regular paragraph
        out.append(f"<p>{line}</p>")
        i += 1

    result = "\n".join(out)
    # Inline styles
    result = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", result)
    result = _re.sub(r"\*(.+?)\*",     r"<em>\1</em>", result)
    result = _re.sub(r"`(.+?)`",       r'<code style="background:#2d2d2d;padding:1px 5px;border-radius:3px;font-family:monospace">\1</code>', result)
    return result


def _build_mime(to: str, subject: str, body_text: str) -> bytes:
    """
    Builds a multipart/alternative MIME message.
    Subject header is UTF-8 encoded via email.header.Header to prevent mojibake.
    If body_text is markdown, an HTML part is also attached.
    """
    import base64 as _b64
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.header import Header

    msg = MIMEMultipart("alternative")
    msg["To"] = to
    msg["Subject"] = Header(subject, "utf-8")

    # Plain text part (fallback)
    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    # HTML part
    html_body = _md_to_html(body_text)
    html_full = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: -apple-system, 'Helvetica Neue', sans-serif;
         background:#0d1117; color:#c9d1d9; padding:24px; max-width:720px; margin:auto; }}
  h1 {{ color:#58a6ff; font-size:20px; margin:16px 0 8px; border-bottom:1px solid #30363d; padding-bottom:6px; }}
  h2 {{ color:#79c0ff; font-size:16px; margin:14px 0 6px; }}
  h3 {{ color:#a5d6ff; font-size:14px; margin:12px 0 4px; }}
  p  {{ margin:4px 0; line-height:1.6; }}
  li {{ margin:2px 0; line-height:1.6; }}
  code {{ background:#2d2d2d; padding:1px 5px; border-radius:3px; font-family:monospace; font-size:13px; }}
  table {{ border-collapse:collapse; width:100%; margin:8px 0; }}
  th {{ background:#161b22; border:1px solid #30363d; padding:6px 10px; text-align:left; color:#79c0ff; }}
  td {{ border:1px solid #21262d; padding:5px 10px; }}
  tr:nth-child(even) td {{ background:#0d1117; }}
  hr {{ border:none; border-top:1px solid #30363d; margin:16px 0; }}
  strong {{ color:#e6edf3; }}
</style></head>
<body>{html_body}</body></html>"""
    msg.attach(MIMEText(html_full, "html", "utf-8"))

    return msg.as_bytes()


def gmail_send(to: str, subject: str, body: str, account: str = "") -> dict:
    import base64
    raw_bytes = _build_mime(to, subject, body)
    encoded = base64.urlsafe_b64encode(raw_bytes).decode()
    base = "gmail.googleapis.com/gmail/v1/users/me"
    return _gapi(f"{base}/messages/send", account=account,
                 method="POST", body={"raw": encoded})


def gmail_draft(to: str, subject: str, body: str, account: str = "") -> dict:
    import base64
    raw_bytes = _build_mime(to, subject, body)
    encoded = base64.urlsafe_b64encode(raw_bytes).decode()
    base = "gmail.googleapis.com/gmail/v1/users/me"
    return _gapi(f"{base}/drafts", account=account,
                 method="POST", body={"message": {"raw": encoded}})


def gmail_modify_labels(message_id: str, add: list[str] | None = None,
                        remove: list[str] | None = None,
                        account: str = "") -> dict:
    base = "gmail.googleapis.com/gmail/v1/users/me"
    return _gapi(f"{base}/messages/{message_id}/modify", account=account,
                 method="POST",
                 body={"addLabelIds": add or [], "removeLabelIds": remove or []})


def gmail_batch_modify(
    message_ids: list[str],
    add: list[str] | None = None,
    remove: list[str] | None = None,
    account: str = "",
) -> dict:
    """
    Modifies labels on multiple messages in a single call using the Gmail batchModify API.
    Reduces API round-trips compared to calling gmail_modify_labels per message.

    Example: mark read + archive = add=[], remove=['UNREAD','INBOX']
    Example: delete newsletter  = add=['TRASH'], remove=['INBOX']
    """
    if not message_ids:
        return {"ok": True, "modified": 0}
    base = "gmail.googleapis.com/gmail/v1/users/me"
    # Gmail batchModify allows up to 1000 IDs per request
    CHUNK = 1000
    total = 0
    for i in range(0, len(message_ids), CHUNK):
        chunk = message_ids[i : i + CHUNK]
        _gapi(
            f"{base}/messages/batchModify",
            account=account,
            method="POST",
            body={
                "ids": chunk,
                "addLabelIds": add or [],
                "removeLabelIds": remove or [],
            },
        )
        total += len(chunk)
    return {"ok": True, "modified": total}


# ── Google Calendar ────────────────────────────────────────────────────────────

def _calendar_ids(account: str) -> list[tuple[str, str]]:
    data = _gapi("www.googleapis.com/calendar/v3/users/me/calendarList", account=account)
    return [(c["id"], c.get("summary", "")) for c in data.get("items", [])]


def calendar_list_events(
    days_from_today: int = 7,
    max_results: int = 20,
    account: str = "",
    calendar_name: str = "",
) -> list[dict]:
    from datetime import timedelta
    now = datetime.now(timezone.utc)
    time_min = now.isoformat()
    time_max = (now + timedelta(days=days_from_today)).isoformat()

    calendars = _calendar_ids(account)
    if calendar_name:
        calendars = [(cid, name) for cid, name in calendars if calendar_name.lower() in name.lower()]

    result = []
    for cal_id, cal_name in calendars:
        try:
            data = _gapi(
                f"www.googleapis.com/calendar/v3/calendars/{urllib.parse.quote(cal_id)}/events",
                account=account,
                params={
                    "timeMin": time_min,
                    "timeMax": time_max,
                    "maxResults": max_results,
                    "singleEvents": "true",
                    "orderBy": "startTime",
                },
            )
        except Exception:
            continue
        for ev in data.get("items", []):
            start = ev.get("start", {})
            end = ev.get("end", {})
            result.append({
                "id": ev.get("id", ""),
                "calendar": cal_name,
                "summary": ev.get("summary", "(제목 없음)"),
                "start": start.get("dateTime", start.get("date", "")),
                "end": end.get("dateTime", end.get("date", "")),
                "location": ev.get("location", ""),
                "description": ev.get("description", "")[:500],
                "link": ev.get("htmlLink", ""),
            })

    result.sort(key=lambda e: e["start"])
    return result


def calendar_create_event(
    summary: str,
    start_iso: str,
    end_iso: str,
    description: str = "",
    location: str = "",
    account: str = "",
) -> dict:
    body = {
        "summary": summary,
        "start": {"dateTime": start_iso, "timeZone": "Asia/Seoul"},
        "end": {"dateTime": end_iso, "timeZone": "Asia/Seoul"},
    }
    if description:
        body["description"] = description
    if location:
        body["location"] = location

    data = _gapi(
        "www.googleapis.com/calendar/v3/calendars/primary/events",
        account=account,
        method="POST",
        body=body,
    )
    return {"id": data.get("id"), "summary": data.get("summary"), "link": data.get("htmlLink")}


def calendar_update_event(
    event_id: str,
    summary: str = "",
    start_iso: str = "",
    end_iso: str = "",
    description: str = "",
    location: str = "",
    account: str = "",
) -> dict:
    body: dict = {}
    if summary:     body["summary"] = summary
    if description: body["description"] = description
    if location:    body["location"] = location
    if start_iso:   body["start"] = {"dateTime": start_iso, "timeZone": "Asia/Seoul"}
    if end_iso:     body["end"]   = {"dateTime": end_iso,   "timeZone": "Asia/Seoul"}
    if not body:
        return {"error": "수정할 내용이 없음"}
    data = _gapi(
        f"www.googleapis.com/calendar/v3/calendars/primary/events/{event_id}",
        account=account,
        method="PATCH",
        body=body,
    )
    return {"id": data.get("id"), "summary": data.get("summary"), "link": data.get("htmlLink")}


def calendar_delete_event(event_id: str, account: str = "") -> dict:
    _gapi(
        f"www.googleapis.com/calendar/v3/calendars/primary/events/{event_id}",
        account=account,
        method="DELETE",
    )
    return {"ok": True, "deleted_id": event_id}


# ── Google Drive ───────────────────────────────────────────────────────────────

def drive_search(query: str, max_results: int = 10, account: str = "") -> list[dict]:
    import re as _re2
    q = query.strip()
    is_drive_syntax = bool(_re2.search(
        r'\b(contains|trashed|mimeType|modifiedTime|name|fullText|parents)\b|[=!<>]', q
    ))
    if not is_drive_syntax:
        safe = q.replace("'", "\\'")
        q = f"fullText contains '{safe}' and trashed = false"

    data = _gapi(
        "www.googleapis.com/drive/v3/files",
        account=account,
        params={
            "q": q,
            "pageSize": max_results,
            "fields": "files(id,name,mimeType,modifiedTime,webViewLink)",
            "orderBy": "modifiedTime desc",
        },
    )
    return [
        {
            "id": f["id"],
            "name": f["name"],
            "mimeType": f["mimeType"],
            "modifiedTime": f.get("modifiedTime", ""),
            "webViewLink": f.get("webViewLink", ""),
        }
        for f in data.get("files", [])
    ]


def _pdf_bytes_to_text(pdf_bytes: bytes, name: str = "") -> str:
    import io
    # pypdf 단독 — requirements.txt·vega-backend.spec에 번들돼 dev/배포 동일하게 동작.
    # (과거 pdfminer 1차 시도는 미번들이라 배포본에선 안 타던 dev-전용 경로 → 제거, INT-1561)
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        return (f"[{name}]\n\n" if name else "") + text[:8000]
    except ImportError:
        return f"[{name}] PDF — pypdf not installed"
    except Exception as e:
        return f"[{name}] PDF 추출 실패: {e}"


def _docx_to_html(path: str) -> str:
    """docx → HTML conversion (uses mammoth). Images are inlined as base64.
    Intended for near-visual Word preview. Falls back to _docx_to_markdown."""
    try:
        import mammoth
    except ImportError:
        return _docx_to_markdown(path)
    try:
        # Style map — 'Title' is not in mammoth's defaults so we define it explicitly
        style_map = """
        p[style-name='Title'] => h1.doc-title:fresh
        p[style-name='Subtitle'] => h2.doc-subtitle:fresh
        p[style-name='Quote'] => blockquote > p:fresh
        p[style-name='Intense Quote'] => blockquote.intense > p:fresh
        """
        with open(path, "rb") as f:
            result = mammoth.convert_to_html(f, style_map=style_map)
        return result.value or "<p><em>(empty document)</em></p>"
    except Exception as e:
        # Fallback to markdown if mammoth fails
        return f"<!-- mammoth failed: {e} -->\n" + _docx_to_markdown(path)


def _docx_to_markdown(path: str) -> str:
    """docx → Markdown conversion. Preserves paragraph styles (Heading), tables, lists, bold/italic.
    Claude Desktop level — layout, images, and pagination are ignored; structure only."""
    try:
        from docx import Document
    except ImportError:
        return "_python-docx not installed_"
    doc = Document(path)
    out: list[str] = []

    def _para_md(para):
        """Inline bold/italic markdown per run."""
        parts = []
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:   t = f"***{t}***"
            elif run.bold:                t = f"**{t}**"
            elif run.italic:              t = f"*{t}*"
            parts.append(t)
        return "".join(parts) or para.text

    # Body paragraphs
    for para in doc.paragraphs:
        if not para.text.strip():
            out.append("")
            continue
        style = (para.style.name or "").lower()
        text = _para_md(para)
        if style.startswith("heading"):
            # 'Heading 1' → '# ', 'Heading 2' → '## ', etc.
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            level = max(1, min(level, 6))
            out.append(f"{'#' * level} {text}")
        elif style.startswith("title"):
            out.append(f"# {text}")
        elif style.startswith("subtitle"):
            out.append(f"## {text}")
        elif style.startswith("list bullet") or "list bullet" in style:
            out.append(f"- {text}")
        elif style.startswith("list number") or "list number" in style:
            out.append(f"1. {text}")
        elif style.startswith("quote"):
            out.append(f"> {text}")
        elif style.startswith("code"):
            out.append(f"```\n{para.text}\n```")
        else:
            out.append(text)

    # Tables — convert to markdown tables and append at the end
    for ti, tbl in enumerate(doc.tables):
        if not tbl.rows:
            continue
        header_cells = [c.text.strip().replace("\n", " ") or " " for c in tbl.rows[0].cells]
        if not header_cells:
            continue
        out.append("")
        out.append("| " + " | ".join(header_cells) + " |")
        out.append("|" + "|".join(["---"] * len(header_cells)) + "|")
        for row in tbl.rows[1:]:
            cells = [c.text.strip().replace("\n", " ").replace("|", "\\|") or " " for c in row.cells]
            # Trim excess cells or pad with empty cells to match header width
            cells = cells[:len(header_cells)] + [" "] * (len(header_cells) - len(cells))
            out.append("| " + " | ".join(cells) + " |")
        out.append("")

    return "\n".join(out).strip() or "_empty document_"


def _xlsx_bytes_to_text(data: bytes, name: str = "", max_rows: int = 500,
                        password: str | None = None) -> str:
    import io
    import openpyxl

    # Auto-detect and decrypt encrypted OLE2 files (.xlsx extension but CDFV2 container)
    if data[:4] == b"\xd0\xcf\x11\xe0":  # OLE2 magic bytes
        try:
            import msoffcrypto
            of = msoffcrypto.OfficeFile(io.BytesIO(data))
            if of.is_encrypted():
                if not password:
                    return ("__needs_password__: 이 엑셀은 비밀번호로 암호화돼 있어. "
                            "file_read를 password 인자와 함께 다시 호출해줘 "
                            f"(파일: {name or '엑셀'}).")
                of.load_key(password=password)
                out = io.BytesIO()
                of.decrypt(out)
                data = out.getvalue()
        except ImportError:
            return "Decryption library (msoffcrypto) not found — run: pip install msoffcrypto-tool"
        except Exception as e:
            return f"Excel decryption failed (check password): {e}"

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        rows = [r for r in rows if any(c is not None for c in r)]
        if not rows:
            continue
        parts.append(f"### Sheet: {sheet_name}")
        header = rows[0]
        data_rows = rows[1:max_rows]
        col_count = max(len(r) for r in rows)
        header_cells = [str(h) if h is not None else "" for h in header]
        if all(h == "" for h in header_cells):
            header_cells = [chr(65 + i) for i in range(col_count)]
            data_rows = rows[:max_rows]
        sep = ["---"] * len(header_cells)
        table_lines = ["| " + " | ".join(header_cells) + " |",
                       "| " + " | ".join(sep) + " |"]
        for row in data_rows:
            cells = [str(c) if c is not None else "" for c in row]
            while len(cells) < len(header_cells):
                cells.append("")
            table_lines.append("| " + " | ".join(cells[:len(header_cells)]) + " |")
        parts.append("\n".join(table_lines))
        if len(rows) - 1 > max_rows:
            parts.append(f"_(showing {max_rows} of {len(rows)-1} rows)_")
    wb.close()
    header_str = f"[{name}]\n\n" if name else ""
    return header_str + "\n\n".join(parts)


def _sqlite_to_text(p: Path, table: str | None = None, max_rows: int = 200) -> str:
    """Reads a SQLite DB and returns schema + sample data as text.
    If table is specified, returns only that table; otherwise summarises all tables."""
    import sqlite3
    try:
        con = sqlite3.connect(f"file:{p}?mode=ro", uri=True)
    except sqlite3.OperationalError as e:
        return f"Failed to open SQLite: {e}"

    try:
        cur = con.cursor()
        # Table list
        cur.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
        all_tables = [r[0] for r in cur.fetchall()]
        if not all_tables:
            return f"[{p.name}] no tables (empty DB)"

        targets = [table] if table and table in all_tables else all_tables
        if table and table not in all_tables:
            return f"Table '{table}' not found. Available tables: {', '.join(all_tables)}"

        parts: list[str] = [f"# {p.name}  ({len(all_tables)} tables)\n"]
        if len(all_tables) > 1 and not table:
            parts.append("Tables: " + ", ".join(all_tables) + "\n")

        char_budget = 7000
        used = sum(len(s) for s in parts)

        for tbl in targets:
            if used >= char_budget:
                parts.append("\n… (character budget reached, remaining tables omitted)")
                break

            # Schema
            cur.execute(f"PRAGMA table_info([{tbl}])")
            cols = cur.fetchall()  # (cid, name, type, notnull, dflt, pk)
            schema_line = ", ".join(
                f"{c[1]} {c[2]}" + (" PK" if c[5] else "") + (" NOT NULL" if c[3] else "")
                for c in cols
            )

            # Row count
            try:
                cur.execute(f"SELECT COUNT(*) FROM [{tbl}]")
                row_count = cur.fetchone()[0]
            except Exception:
                row_count = "?"

            header = f"\n## {tbl}  ({row_count} rows)\nSchema: {schema_line}\n"
            parts.append(header)
            used += len(header)

            # Sample data
            limit = min(max_rows, 50)
            try:
                cur.execute(f"SELECT * FROM [{tbl}] LIMIT {limit}")
                rows = cur.fetchall()
                col_names = [c[1] for c in cols]
                rows_text = "\t".join(col_names) + "\n"
                for row in rows:
                    rows_text += "\t".join("" if v is None else str(v) for v in row) + "\n"
                if row_count != "?" and row_count > limit:
                    rows_text += f"… ({row_count - limit} more rows)\n"
            except Exception as e:
                rows_text = f"(failed to read data: {e})\n"

            parts.append(rows_text)
            used += len(rows_text)

        return "".join(parts)
    finally:
        con.close()


def _local_excel_to_text(path: str, sheet: str | None = None, max_rows: int = 500,
                         password: str | None = None) -> str:
    p = Path(path).expanduser()
    if not p.exists():
        return f"파일 없음: {path}"
    suffix = p.suffix.lower()
    if suffix in (".csv", ".tsv"):
        sep = "\t" if suffix == ".tsv" else ","
        import csv
        with open(p, newline="", encoding="utf-8-sig", errors="replace") as f:
            reader = csv.reader(f, delimiter=sep)
            rows = [r for _, r in zip(range(max_rows + 1), reader)]
        if not rows:
            return f"[{p.name}] empty file"
        header = rows[0]
        sep_row = ["---"] * len(header)
        lines = ["| " + " | ".join(header) + " |",
                 "| " + " | ".join(sep_row) + " |"]
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[:len(header)]) + " |")
        return f"[{p.name}]\n\n" + "\n".join(lines)
    elif suffix in (".xlsx", ".xls", ".xlsm"):
        return _xlsx_bytes_to_text(p.read_bytes(), p.name, max_rows=max_rows, password=password)
    else:
        return f"Unsupported format: {suffix} (supported: xlsx, xls, xlsm, csv, tsv)"


def file_edit(
    path: str,
    operation: str,
    sheet: str | None = None,
    row: int | None = None,
    col: int | None = None,
    col_name: str | None = None,
    value: Any = None,
    values: list | None = None,
    where: dict | None = None,
) -> str:
    import shutil
    p = Path(path).expanduser()
    if not p.exists():
        return f"파일 없음: {path}"
    suffix = p.suffix.lower()
    bak = p.with_suffix(p.suffix + ".bak")
    shutil.copy2(p, bak)
    try:
        if suffix in (".xlsx", ".xlsm"):
            return _xlsx_edit(p, operation, sheet, row, col, col_name, value, values, where)
        elif suffix in (".csv", ".tsv"):
            return _csv_edit(p, suffix, operation, row, col, col_name, value, values, where)
        else:
            return f"Cannot edit: {suffix} (supported: xlsx, xlsm, csv, tsv)"
    except Exception as e:
        shutil.copy2(bak, p)
        return f"Edit failed (file restored from backup): {e}"


def _xlsx_edit(
    p: Path, operation: str,
    sheet: str | None, row: int | None, col: int | None,
    col_name: str | None, value: Any, values: list | None, where: dict | None,
) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(p)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active

    if operation == "set_cell":
        if col_name:
            headers = [str(ws.cell(1, c).value) for c in range(1, ws.max_column + 1)]
            try:
                col = headers.index(col_name) + 1
            except ValueError:
                return f"Column name not found: {col_name}"
        if row is None or col is None:
            return "set_cell: row and col (or col_name) are required"
        ws.cell(row=row, column=col, value=value)
        wb.save(p)
        return f"✓ [{ws.title}] ({row},{col}) = {value!r}"
    elif operation == "append_row":
        if not values:
            return "append_row: values is required"
        ws.append(values)
        wb.save(p)
        return f"✓ [{ws.title}] row appended: {values}"
    elif operation == "update_row":
        if not where or not values:
            return "update_row: where and values are required"
        headers = [str(ws.cell(1, c).value) for c in range(1, ws.max_column + 1)]
        updated = 0
        for r in range(2, ws.max_row + 1):
            row_data = {headers[c]: ws.cell(r, c + 1).value for c in range(len(headers))}
            if all(str(row_data.get(k, "")) == str(v) for k, v in where.items()):
                for k, v in values.items():
                    if k in headers:
                        ws.cell(r, headers.index(k) + 1, value=v)
                updated += 1
        wb.save(p)
        return f"✓ {updated} rows updated"
    elif operation == "delete_row":
        if not where:
            return "delete_row: where is required"
        headers = [str(ws.cell(1, c).value) for c in range(1, ws.max_column + 1)]
        to_delete = []
        for r in range(2, ws.max_row + 1):
            row_data = {headers[c]: ws.cell(r, c + 1).value for c in range(len(headers))}
            if all(str(row_data.get(k, "")) == str(v) for k, v in where.items()):
                to_delete.append(r)
        for r in reversed(to_delete):
            ws.delete_rows(r)
        wb.save(p)
        return f"✓ {len(to_delete)} rows deleted"
    elif operation == "add_sheet":
        if not sheet:
            return "add_sheet: sheet name is required"
        if sheet in wb.sheetnames:
            return f"Sheet already exists: {sheet}"
        wb.create_sheet(sheet)
        wb.save(p)
        return f"✓ sheet added: {sheet}"
    else:
        return f"Unknown operation: {operation}"


def _csv_edit(
    p: Path, suffix: str, operation: str,
    row: int | None, col: int | None, col_name: str | None,
    value: Any, values: list | None, where: dict | None,
) -> str:
    import csv
    delim = "\t" if suffix == ".tsv" else ","
    with open(p, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.reader(f, delimiter=delim))
    if not rows:
        return "empty file"
    headers = rows[0]

    if operation == "set_cell":
        if col_name:
            try:
                col = headers.index(col_name) + 1
            except ValueError:
                return f"Column name not found: {col_name}"
        if row is None or col is None:
            return "set_cell: row and col (or col_name) are required"
        while len(rows) < row:
            rows.append([""] * len(headers))
        while len(rows[row - 1]) < col:
            rows[row - 1].append("")
        rows[row - 1][col - 1] = str(value) if value is not None else ""
        msg = f"✓ ({row},{col}) = {value!r}"
    elif operation == "append_row":
        if not values:
            return "append_row: values is required"
        rows.append([str(v) if v is not None else "" for v in values])
        msg = f"✓ row appended: {values}"
    elif operation == "update_row":
        if not where or not values:
            return "update_row: where and values are required"
        updated = 0
        for r in rows[1:]:
            row_dict = dict(zip(headers, r))
            if all(str(row_dict.get(k, "")) == str(v) for k, v in where.items()):
                for k, v in values.items():
                    if k in headers:
                        r[headers.index(k)] = str(v) if v is not None else ""
                updated += 1
        msg = f"✓ {updated} rows updated"
    elif operation == "delete_row":
        if not where:
            return "delete_row: where is required"
        before = len(rows)
        rows = [rows[0]] + [
            r for r in rows[1:]
            if not all(str(dict(zip(headers, r)).get(k, "")) == str(v) for k, v in where.items())
        ]
        msg = f"✓ {before - len(rows)} rows deleted"
    else:
        return f"Unknown operation: {operation} (csv supports: set_cell, append_row, update_row, delete_row)"

    with open(p, "w", newline="", encoding="utf-8-sig") as f:
        csv.writer(f, delimiter=delim).writerows(rows)
    return msg


_ICLOUD_ROOT = Path.home() / "Library/Mobile Documents/com~apple~CloudDocs"


def _resolve_icloud_path(path: str) -> Path:
    for prefix in ("~/iCloud/", "~/iCloud Drive/", "iCloud:", "icloud:"):
        if path.startswith(prefix):
            rel = path[len(prefix):]
            return _ICLOUD_ROOT / rel
    p = Path(path).expanduser()
    if not p.is_absolute():
        return _ICLOUD_ROOT / path
    return p


def icloud_list(path: str = "") -> list[dict]:
    target = _ICLOUD_ROOT / path if path else _ICLOUD_ROOT
    if not target.exists():
        return [{"error": f"경로 없음: {target}"}]
    items = []
    for item in sorted(i for i in target.iterdir() if not i.name.startswith(".")):
        stat = item.stat()
        items.append({
            "name": item.name,
            "type": "dir" if item.is_dir() else "file",
            "size": stat.st_size if item.is_file() else None,
            "path": str(item),
        })
    return items


def icloud_move(src: str, dst: str) -> dict:
    s = _resolve_icloud_path(src)
    d = _resolve_icloud_path(dst)
    if not s.exists():
        return {"ok": False, "error": f"소스 없음: {s}"}
    if d.is_dir():
        d = d / s.name
    if d.exists():
        return {"ok": False, "error": f"대상 이미 존재: {d}"}
    d.parent.mkdir(parents=True, exist_ok=True)
    s.rename(d)
    return {"ok": True, "src": str(s), "dst": str(d)}


def icloud_rename(path: str, new_name: str) -> dict:
    p = _resolve_icloud_path(path)
    if not p.exists():
        return {"ok": False, "error": f"경로 없음: {p}"}
    if "/" in new_name or "\\" in new_name:
        return {"ok": False, "error": "new_name에 경로 구분자 불가"}
    dest = p.parent / new_name
    if dest.exists():
        return {"ok": False, "error": f"이미 존재: {dest}"}
    p.rename(dest)
    return {"ok": True, "old": str(p), "new": str(dest)}


def icloud_mkdir(path: str) -> dict:
    p = _resolve_icloud_path(path)
    if p.exists():
        return {"ok": False, "error": f"이미 존재: {p}"}
    p.mkdir(parents=True)
    return {"ok": True, "path": str(p)}


def file_read(path: str, sheet: str | None = None, max_rows: int = 500,
              password: str | None = None) -> list[dict] | str:
    if "iCloud" in path or "icloud" in path:
        p = _resolve_icloud_path(path)
    else:
        p = Path(path).expanduser()
        # Relative paths are resolved relative to the current session working directory
        if not p.is_absolute():
            from pipeline.tools_code import _base_cwd
            p = Path(_base_cwd()) / p
    # Block sensitive paths (checked after symlink resolution)
    try:
        from pipeline.path_guard import guard_path
        p = guard_path(str(p))
    except PermissionError as e:
        return f"[SAFEGUARD] 접근 차단: {e}"
    if not p.exists():
        return f"파일 없음: {path}"
    suffix = p.suffix.lower()
    if suffix in (".xlsx", ".xls", ".xlsm", ".csv", ".tsv"):
        return _local_excel_to_text(str(p), sheet=sheet, max_rows=max_rows, password=password)
    elif suffix in (".db", ".sqlite", ".sqlite3"):
        return _sqlite_to_text(p, table=sheet, max_rows=max_rows)
    elif suffix in (".txt", ".md", ".json", ".py", ".yaml", ".yml", ".toml"):
        return p.read_text(errors="replace")[:8000]
    else:
        try:
            return p.read_text(errors="replace")[:8000]
        except Exception as e:
            return f"Read failed: {e}"


def drive_read(file_id: str, account: str = "") -> str:
    import re as _re
    url_match = _re.search(r'[?&]id=([a-zA-Z0-9_-]+)', file_id)
    if url_match:
        file_id = url_match.group(1)

    try:
        meta = _gapi(
            f"www.googleapis.com/drive/v3/files/{file_id}",
            account=account,
            params={"fields": "id,name,mimeType"},
        )
    except Exception:
        meta = None
    if meta is None:
        return _drive_shared_download(file_id)

    mime = meta.get("mimeType", "")
    name = meta.get("name", file_id)

    export_map = {
        "application/vnd.google-apps.document": "text/plain",
        "application/vnd.google-apps.spreadsheet": "text/csv",
        "application/vnd.google-apps.presentation": "text/plain",
    }
    token = _google_token(account)
    export_mime = export_map.get(mime)

    if export_mime:
        url = (f"https://www.googleapis.com/drive/v3/files/{file_id}/export"
               f"?mimeType={urllib.parse.quote(export_mime)}")
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        with urllib.request.urlopen(req, timeout=20) as r:
            return r.read().decode("utf-8", errors="replace")[:8000]
    elif mime == "application/pdf":
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        with urllib.request.urlopen(req, timeout=30) as r:
            pdf_bytes = r.read()
        return _pdf_bytes_to_text(pdf_bytes, name)
    elif mime == "text/plain" or mime.startswith("text/"):
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        with urllib.request.urlopen(req, timeout=20) as r:
            return r.read().decode("utf-8", errors="replace")[:8000]
    elif mime in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
        req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
        with urllib.request.urlopen(req, timeout=30) as r:
            xlsx_bytes = r.read()
        return _xlsx_bytes_to_text(xlsx_bytes, name)
    else:
        return f"[{name}] ({mime}) — cannot convert to text (supported: Docs, Sheets, Slides, PDF, xlsx, plain text)"


def _drive_shared_download(file_id: str) -> str:
    import concurrent.futures, tempfile, os, asyncio

    def _worker():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            from playwright.sync_api import sync_playwright
            download_url = f"https://drive.google.com/uc?export=download&id={file_id}&confirm=t"
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                context = browser.new_context(accept_downloads=True)
                page = context.new_page()
                with tempfile.TemporaryDirectory() as tmpdir:
                    with page.expect_download(timeout=30_000) as dl_info:
                        try:
                            page.goto(download_url, timeout=20_000, wait_until="commit")
                        except Exception:
                            pass
                    download = dl_info.value
                    dest = os.path.join(tmpdir, download.suggested_filename or "file")
                    download.save_as(dest)
                    browser.close()
                    fname = os.path.basename(dest)
                    if dest.lower().endswith(".pdf"):
                        with open(dest, "rb") as f:
                            return _pdf_bytes_to_text(f.read(), fname)
                    else:
                        try:
                            with open(dest, "r", encoding="utf-8", errors="replace") as f:
                                return f"[{fname}]\n\n" + f.read()[:8000]
                        except Exception:
                            return f"[{fname}] binary file — cannot convert to text"
        finally:
            loop.close()

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(_worker)
            return future.result(timeout=90)
    except Exception as e:
        return f"Drive shared file download failed: {e}"


# ── Google Slides ──────────────────────────────────────────────────────────────

def _slides_batch_update(presentation_id: str, requests: list[dict], account: str) -> dict:
    return _gapi(
        f"slides.googleapis.com/v1/presentations/{presentation_id}:batchUpdate",
        account=account,
        method="POST",
        body={"requests": requests},
    )


def _slide_requests(slides: list[dict], start_index: int = 0) -> list[dict]:
    """Converts a list of slide specs into Slides API batchUpdate requests.

    Slide spec format:
      title   : title text
      body    : body text (newlines supported)
      layout  : 'TITLE_AND_BODY' (default) | 'TITLE_ONLY' | 'BLANK' | 'SECTION_HEADER'
      notes   : presenter notes
    """
    import uuid
    reqs: list[dict] = []

    layout_map = {
        "TITLE_AND_BODY":  "TITLE_AND_BODY",
        "TITLE_ONLY":      "TITLE_ONLY",
        "BLANK":           "BLANK",
        "SECTION_HEADER":  "SECTION_HEADER",
        "title_and_body":  "TITLE_AND_BODY",
        "title_only":      "TITLE_ONLY",
        "blank":           "BLANK",
        "section":         "SECTION_HEADER",
    }

    for i, spec in enumerate(slides):
        slide_id = f"slide_{start_index + i}_{uuid.uuid4().hex[:6]}"
        layout = layout_map.get(spec.get("layout", ""), "TITLE_AND_BODY")

        reqs.append({
            "createSlide": {
                "objectId": slide_id,
                "insertionIndex": start_index + i,
                "slideLayoutReference": {"predefinedLayout": layout},
            }
        })

        # Title
        title_text = spec.get("title", "")
        if title_text:
            title_id = f"{slide_id}_title"
            reqs.append({
                "createShape": {
                    "objectId": title_id,
                    "shapeType": "TEXT_BOX",
                    "elementProperties": {
                        "pageObjectId": slide_id,
                        "size": {"width": {"magnitude": 6096000, "unit": "EMU"},
                                 "height": {"magnitude": 900000, "unit": "EMU"}},
                        "transform": {"scaleX": 1, "scaleY": 1,
                                      "translateX": 457200, "translateY": 274638,
                                      "unit": "EMU"},
                    },
                }
            })
            reqs.append({
                "insertText": {
                    "objectId": title_id,
                    "insertionIndex": 0,
                    "text": title_text,
                }
            })
            reqs.append({
                "updateTextStyle": {
                    "objectId": title_id,
                    "style": {"bold": True, "fontSize": {"magnitude": 28, "unit": "PT"}},
                    "textRange": {"type": "ALL"},
                    "fields": "bold,fontSize",
                }
            })

        # Body
        body_text = spec.get("body", "")
        if body_text:
            body_id = f"{slide_id}_body"
            reqs.append({
                "createShape": {
                    "objectId": body_id,
                    "shapeType": "TEXT_BOX",
                    "elementProperties": {
                        "pageObjectId": slide_id,
                        "size": {"width": {"magnitude": 6096000, "unit": "EMU"},
                                 "height": {"magnitude": 3500000, "unit": "EMU"}},
                        "transform": {"scaleX": 1, "scaleY": 1,
                                      "translateX": 457200, "translateY": 1310000,
                                      "unit": "EMU"},
                    },
                }
            })
            reqs.append({
                "insertText": {
                    "objectId": body_id,
                    "insertionIndex": 0,
                    "text": body_text,
                }
            })
            reqs.append({
                "updateTextStyle": {
                    "objectId": body_id,
                    "style": {"fontSize": {"magnitude": 18, "unit": "PT"}},
                    "textRange": {"type": "ALL"},
                    "fields": "fontSize",
                }
            })

        # Presenter notes
        notes_text = spec.get("notes", "")
        if notes_text:
            notes_page_id = f"{slide_id}_notes"
            reqs.append({
                "createShape": {
                    "objectId": notes_page_id,
                    "shapeType": "TEXT_BOX",
                    "elementProperties": {
                        "pageObjectId": f"{slide_id}_notes_page",
                        "size": {"width": {"magnitude": 6096000, "unit": "EMU"},
                                 "height": {"magnitude": 2000000, "unit": "EMU"}},
                        "transform": {"scaleX": 1, "scaleY": 1,
                                      "translateX": 457200, "translateY": 457200,
                                      "unit": "EMU"},
                    },
                }
            })
            reqs.append({
                "insertText": {
                    "objectId": notes_page_id,
                    "insertionIndex": 0,
                    "text": notes_text,
                }
            })

    return reqs


def slides_create(
    title: str,
    slides: list[dict],
    account: str = "",
) -> dict:
    """Creates a new Google Slides presentation and populates it with slides.

    Slide item format:
      {"title": "Title", "body": "Body\\n- item1\\n- item2",
       "layout": "TITLE_AND_BODY", "notes": "presenter notes"}

    Returns: {"ok": True, "presentation_id": "...", "url": "...", "slides_created": N}
    """
    # Create an empty presentation
    pres = _gapi(
        "slides.googleapis.com/v1/presentations",
        account=account,
        method="POST",
        body={"title": title},
    )
    pid = pres["presentationId"]

    # Delete the default empty slide before adding the requested slides
    existing = pres.get("slides", [])
    reqs: list[dict] = []
    for s in existing:
        reqs.append({"deleteObject": {"objectId": s["objectId"]}})
    reqs.extend(_slide_requests(slides, start_index=0))

    if reqs:
        _slides_batch_update(pid, reqs, account)

    return {
        "ok": True,
        "presentation_id": pid,
        "url": f"https://docs.google.com/presentation/d/{pid}/edit",
        "slides_created": len(slides),
    }


def slides_append_slide(
    presentation_id: str,
    slides: list[dict],
    account: str = "",
) -> dict:
    """Appends slides to an existing Google Slides presentation."""
    pres = _gapi(
        f"slides.googleapis.com/v1/presentations/{presentation_id}",
        account=account,
    )
    current_count = len(pres.get("slides", []))
    reqs = _slide_requests(slides, start_index=current_count)
    if reqs:
        _slides_batch_update(presentation_id, reqs, account)
    return {
        "ok": True,
        "presentation_id": presentation_id,
        "url": f"https://docs.google.com/presentation/d/{presentation_id}/edit",
        "slides_added": len(slides),
        "total_slides": current_count + len(slides),
    }


# ── Google Docs ────────────────────────────────────────────────────────────────

def _docs_batch_update(document_id: str, requests: list[dict], account: str) -> dict:
    return _gapi(
        f"docs.googleapis.com/v1/documents/{document_id}:batchUpdate",
        account=account,
        method="POST",
        body={"requests": requests},
    )


def _content_to_doc_requests(content: list[dict], insert_at: int = 1) -> list[dict]:
    """Converts a list of content blocks into Docs API batchUpdate requests.

    Block format:
      {"type": "heading", "text": "Title", "level": 1}   # level 1–6
      {"type": "paragraph", "text": "body text"}
      {"type": "bullet", "text": "item"}                  # bullet list
      {"type": "table", "rows": [["A","B"],["1","2"]]}
      {"type": "pagebreak"}
    """
    reqs: list[dict] = []
    # The Docs API requires insertText location.index to be accumulated after each insertion.
    # Simplest approach: insert in order and track the running index.
    # We insert sequentially at the end of the document (index=insert_at).
    idx = insert_at

    heading_style_map = {1: "HEADING_1", 2: "HEADING_2", 3: "HEADING_3",
                         4: "HEADING_4", 5: "HEADING_5", 6: "HEADING_6"}

    for block in content:
        btype = block.get("type", "paragraph")
        text = block.get("text", "")

        if btype == "pagebreak":
            reqs.append({"insertPageBreak": {"location": {"index": idx}}})
            idx += 1
            continue

        if btype == "table":
            rows = block.get("rows", [])
            if not rows:
                continue
            n_rows = len(rows)
            n_cols = max(len(r) for r in rows)
            reqs.append({
                "insertTable": {
                    "rows": n_rows,
                    "columns": n_cols,
                    "location": {"index": idx},
                }
            })
            # Cell content requires separate requests after table insertion — position math is complex.
            # Fallback: append each row as a tab-separated paragraph below the table (API limitation).
            idx += 1
            for row in rows:
                row_text = "\t".join(str(v) for v in row) + "\n"
                reqs.append({"insertText": {"location": {"index": idx}, "text": row_text}})
                idx += len(row_text)
            continue

        # Text block
        line = text + "\n"
        reqs.append({"insertText": {"location": {"index": idx}, "text": line}})

        # Apply style
        end_idx = idx + len(line)
        if btype == "heading":
            level = block.get("level", 1)
            style = heading_style_map.get(level, "HEADING_1")
            reqs.append({
                "updateParagraphStyle": {
                    "range": {"startIndex": idx, "endIndex": end_idx},
                    "paragraphStyle": {"namedStyleType": style},
                    "fields": "namedStyleType",
                }
            })
        elif btype == "bullet":
            reqs.append({
                "createParagraphBullets": {
                    "range": {"startIndex": idx, "endIndex": end_idx},
                    "bulletPreset": "BULLET_DISC_CIRCLE_SQUARE",
                }
            })

        idx = end_idx

    return reqs


def docs_create(
    title: str,
    content: list[dict],
    account: str = "",
) -> dict:
    """Creates a new Google Docs document and inserts content.

    Content block format:
      {"type": "heading", "text": "Title", "level": 1}
      {"type": "paragraph", "text": "body text"}
      {"type": "bullet", "text": "bullet item"}
      {"type": "table", "rows": [["col1","col2"], ["val1","val2"]]}
      {"type": "pagebreak"}

    Returns: {"ok": True, "document_id": "...", "url": "..."}
    """
    doc = _gapi(
        "docs.googleapis.com/v1/documents",
        account=account,
        method="POST",
        body={"title": title},
    )
    did = doc["documentId"]

    if content:
        reqs = _content_to_doc_requests(content, insert_at=1)
        if reqs:
            _docs_batch_update(did, reqs, account)

    return {
        "ok": True,
        "document_id": did,
        "url": f"https://docs.google.com/document/d/{did}/edit",
    }


def docs_append(
    document_id: str,
    content: list[dict],
    account: str = "",
) -> dict:
    """Appends content to the end of an existing Google Docs document."""
    doc = _gapi(
        f"docs.googleapis.com/v1/documents/{document_id}",
        account=account,
    )
    # Current end index of the document
    body = doc.get("body", {})
    content_list = body.get("content", [])
    end_index = content_list[-1].get("endIndex", 1) - 1 if content_list else 1

    reqs = _content_to_doc_requests(content, insert_at=end_index)
    if reqs:
        _docs_batch_update(document_id, reqs, account)

    return {
        "ok": True,
        "document_id": document_id,
        "url": f"https://docs.google.com/document/d/{document_id}/edit",
    }
